"""
agent/parser.py — shared LLM parsing utilities with JSON mode + retry logic.
Used by both graph.py (in-process) and api/main.py (HTTP endpoints).
"""

import io
import json
import logging
import re

import pdfplumber
import docx
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage

from config.settings import settings

logger = logging.getLogger(__name__)


class ParsingError(Exception):
    """Raised when LLM parsing fails after all retries."""
    pass


def _get_fallback_model_details():
    provider = settings.llm_provider
    base_url = settings.resolved_base_url
    api_key = settings.llm_api_key

    if provider == "gemini":
        model = "gemini-2.5-flash"
    elif provider == "groq":
        model = "llama-3.1-70b-versatile"
    elif provider == "openai":
        model = "gpt-4o-mini"
    elif provider == "qwen":
        model = "qwen-plus"
    elif provider == "openrouter":
        model = settings.resolved_model
    else:
        model = settings.resolved_model

    return model, base_url, api_key


def clean_json_str(text: str) -> str:
    text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()
    if text.startswith("```"):
        lines = text.split("\n")
        text = "\n".join(lines[1:-1]) if lines[-1].strip() == "```" else "\n".join(lines[1:])
    text = text.strip()
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return text[start:end+1]
    return text


def _call_llm_json(system_prompt: str, human_prompt: str, response_schema) -> dict:
    """
    Calls LLM in JSON mode, validates against the response_schema.
    Retries with corrective prompt, then falls back to provider-specific model.
    Raises ParsingError if all attempts fail.
    """
    model_name = settings.resolved_model
    api_key = settings.llm_api_key
    base_url = settings.resolved_base_url

    def run_inference(model: str, url: str, key: str, msgs) -> str:
        llm = ChatOpenAI(
            model=model,
            api_key=key,
            base_url=url,
            temperature=0.2,
            max_tokens=4096,
            timeout=120,
        ).bind(response_format={"type": "json_object"})
        resp = llm.invoke(msgs)
        return resp.content.strip()

    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=human_prompt)
    ]

    logger.info(f"Attempting parsing with model={model_name}...")
    try:
        raw_response = run_inference(model_name, base_url, api_key, messages)
        cleaned_json = clean_json_str(raw_response)
        parsed_dict = json.loads(cleaned_json)
        validated_data = response_schema.model_validate(parsed_dict)
        return validated_data.model_dump()
    except Exception as e:
        logger.warning(f"First attempt failed: {str(e)}. Retrying with corrective prompt...")
        bad_response_str = ""
        try:
            bad_response_str = raw_response
        except NameError:
            pass

        corrective_msgs = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        if bad_response_str:
            corrective_msgs.append(AIMessage(content=bad_response_str))

        validation_error_details = str(e)
        feedback_prompt = (
            f"Your previous response was invalid JSON or did not match the required schema.\n"
            f"Validation Error: {validation_error_details}\n"
            f"Please correct the output. Return ONLY the valid JSON object conforming strictly to the schema, "
            f"without any explanation, prefix, suffix, or markdown formatting."
        )
        corrective_msgs.append(HumanMessage(content=feedback_prompt))

        try:
            raw_response2 = run_inference(model_name, base_url, api_key, corrective_msgs)
            cleaned_json2 = clean_json_str(raw_response2)
            parsed_dict2 = json.loads(cleaned_json2)
            validated_data2 = response_schema.model_validate(parsed_dict2)
            return validated_data2.model_dump()
        except Exception as retry_err:
            logger.warning(f"Same-model corrective retry failed: {str(retry_err)}. Falling back to provider fallback model...")
            fb_model, fb_url, fb_key = _get_fallback_model_details()
            logger.info(f"Using fallback model={fb_model}...")

            try:
                raw_response3 = run_inference(fb_model, fb_url, fb_key, corrective_msgs)
                cleaned_json3 = clean_json_str(raw_response3)
                parsed_dict3 = json.loads(cleaned_json3)
                validated_data3 = response_schema.model_validate(parsed_dict3)
                return validated_data3.model_dump()
            except Exception as fb_err:
                logger.error(f"All parsing attempts failed: {str(fb_err)}")
                raise ParsingError(f"Failed to parse after corrective retry and fallback: {str(fb_err)}")


def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """Extracts plain text from PDF, DOCX, or text files."""
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        text_parts = []
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            logger.error(f"Error reading PDF {filename}: {e}")
            raise ValueError(f"Failed to read PDF file: {str(e)}")
        return "\n\n".join(text_parts) if text_parts else ""
    elif filename_lower.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
        except Exception as e:
            logger.error(f"Error reading DOCX {filename}: {e}")
            raise ValueError(f"Failed to read DOCX file: {str(e)}")
        return text
    else:
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error(f"Error decoding text file {filename}: {e}")
            raise ValueError(f"Failed to decode text file: {str(e)}")
