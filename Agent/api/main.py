# ============================================================
# api/main.py — FastAPI app exposing the LangGraph agent
# ============================================================

import asyncio
import logging
import io
import json
from typing import Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import Response, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel, ValidationError
import pdfplumber
import docx

from agent.graph import resume_graph
from agent.state import ResumeAgentState
from tools.resume_tools import check_compiler_health
from config.settings import settings
from api.schemas import JDStructuredOutput, ResumeStructuredOutput
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage, AIMessage

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="ResumeAI Agent",
    description="Agentic resume tailoring: JD → AI → LaTeX → PDF",
    version="1.0.0",
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ================================================================
# REQUEST / RESPONSE MODELS
# ================================================================

class TextResumeRequest(BaseModel):
    candidate_text: str          # raw text describing candidate
    job_description: str
    stream_logs: bool = False


class ResumeResponse(BaseModel):
    success: bool
    message: str
    latex_source: Optional[str] = None
    jd_role_title: Optional[str] = None
    jd_company: Optional[str] = None
    keywords: Optional[list[str]] = None
    compile_attempts: int = 0
    logs: list[str] = []


# ================================================================
# HELPER
# ================================================================

def _run_agent(raw_input: str, job_description: str, input_mode: str) -> ResumeAgentState:
    initial_state: ResumeAgentState = {
        "raw_input": raw_input,
        "job_description": job_description,
        "input_mode": input_mode,
        "jd_keywords": [],
        "jd_role_title": "",
        "jd_company": "",
        "candidate_summary": "",
        "latex_source": "",
        "latex_version": 0,
        "compile_attempts": 0,
        "compile_success": False,
        "compile_error": "",
        "pdf_bytes": None,
        "error": "",
        "status": "starting",
        "logs": [],
    }
    final_state = resume_graph.invoke(
        initial_state,
        config={"recursion_limit": settings.max_agent_iterations}
    )
    return final_state


# ================================================================
# ROUTES
# ================================================================

@app.get("/health")
async def health():
    """Health check for this service + the compiler service."""
    compiler_status = check_compiler_health.invoke({})
    return {
        "agent_api": "ok",
        "compiler": compiler_status,
        "llm_provider": settings.llm_provider,
        "llm_model": settings.resolved_model,
    }


@app.post("/generate/from-text")
async def generate_from_text(req: TextResumeRequest):
    """
    Generate a tailored PDF resume from raw text input.
    Returns the PDF file directly.
    """
    try:
        state = await asyncio.get_event_loop().run_in_executor(
            None, _run_agent, req.candidate_text, req.job_description, "text"
        )
    except Exception as e:
        raise HTTPException(500, f"Agent error: {str(e)}")

    if not state["compile_success"] or not state["pdf_bytes"]:
        raise HTTPException(422, {
            "error": "Resume generation failed",
            "last_error": state.get("compile_error", state.get("error", "")),
            "logs": state.get("logs", []),
        })

    return Response(
        content=state["pdf_bytes"],
        media_type="application/pdf",
        headers={
            "Content-Disposition": "attachment; filename=resume.pdf",
            "X-Latex-Version": str(state.get("latex_version", 1)),
            "X-Compile-Attempts": str(state.get("compile_attempts", 1)),
            "X-JD-Role": state.get("jd_role_title", ""),
        }
    )


@app.post("/generate/from-pdf")
async def generate_from_pdf(
    pdf_file: UploadFile = File(...),
    job_description: str = Form(...),
):
    """
    Upload an existing resume PDF + paste a JD.
    Returns a new tailored PDF resume.
    """
    if not pdf_file.filename.endswith(".pdf"):
        raise HTTPException(400, "Only PDF files accepted")

    pdf_bytes = await pdf_file.read()
    pdf_hex = pdf_bytes.hex()   # pass as hex string to tool

    try:
        state = await asyncio.get_event_loop().run_in_executor(
            None, _run_agent, pdf_hex, job_description, "pdf"
        )
    except Exception as e:
        raise HTTPException(500, f"Agent error: {str(e)}")

    if not state["compile_success"] or not state["pdf_bytes"]:
        raise HTTPException(422, {
            "error": "Resume generation failed",
            "last_error": state.get("compile_error", state.get("error", "")),
            "logs": state.get("logs", []),
        })

    return Response(
        content=state["pdf_bytes"],
        media_type="application/pdf",
        headers={
            "Content-Disposition": "attachment; filename=resume_tailored.pdf",
            "X-Compile-Attempts": str(state.get("compile_attempts", 1)),
            "X-JD-Role": state.get("jd_role_title", ""),
        }
    )


@app.post("/generate/latex-only", response_model=ResumeResponse)
async def generate_latex_only(req: TextResumeRequest):
    """
    Returns the LaTeX source + metadata without compiling.
    Useful for previewing / editing before PDF generation.
    """
    try:
        state = await asyncio.get_event_loop().run_in_executor(
            None, _run_agent, req.candidate_text, req.job_description, "text"
        )
    except Exception as e:
        raise HTTPException(500, str(e))

    return ResumeResponse(
        success=state.get("compile_success", False),
        message="LaTeX generated" if state.get("latex_source") else "Generation failed",
        latex_source=state.get("latex_source"),
        jd_role_title=state.get("jd_role_title"),
        jd_company=state.get("jd_company"),
        keywords=state.get("jd_keywords"),
        compile_attempts=state.get("compile_attempts", 0),
        logs=state.get("logs", []),
    )


@app.post("/compile/raw")
async def compile_raw_latex(latex: str = Form(...)):
    """
    Directly compile raw LaTeX source (bypass the agent).
    Useful for manual edits / re-compile after tweaking.
    """
    from tools.resume_tools import compile_latex
    result = compile_latex.invoke({"latex_source": latex})
    if result["success"]:
        return Response(
            content=bytes.fromhex(result["pdf_hex"]),
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=resume.pdf"}
        )
    raise HTTPException(422, {"error": result["error"]})


# ================================================================
# JD & RESUME PARSING ENDPOINTS & HELPERS
# ================================================================

class ParseJDRequest(BaseModel):
    jd_text: str


def _extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extracts plain text from PDF, DOCX, or text files.
    """
    filename_lower = filename.lower()
    if filename_lower.endswith(".pdf"):
        text_parts = []
        try:
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
        except Exception as e:
            logger.error(f"Error reading PDF {filename}: {e}")
            raise HTTPException(status_code=422, detail=f"Failed to read PDF file: {str(e)}")
        return "\n\n".join(text_parts) if text_parts else ""
    elif filename_lower.endswith(".docx"):
        try:
            doc = docx.Document(io.BytesIO(file_bytes))
            text = "\n".join([p.text for p in doc.paragraphs])
            # Also extract table text if present
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
        except Exception as e:
            logger.error(f"Error reading DOCX {filename}: {e}")
            raise HTTPException(status_code=422, detail=f"Failed to read DOCX file: {str(e)}")
        return text
    else:
        # Fallback to plain text
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception as e:
            logger.error(f"Error decoding text file {filename}: {e}")
            raise HTTPException(status_code=422, detail=f"Failed to decode text file: {str(e)}")


def _get_fallback_model_details():
    provider = settings.llm_provider
    base_url = settings.resolved_base_url
    api_key = settings.llm_api_key
    
    # We choose a highly reliable fallback model supported by the provider
    if provider == "gemini":
        model = "gemini-2.5-flash"
    elif provider == "groq":
        model = "llama-3.1-70b-versatile"
    elif provider == "openai":
        model = "gpt-4o-mini"
    elif provider == "qwen":
        model = "qwen-plus"
    else:
        model = settings.resolved_model
        
    return model, base_url, api_key


def _call_llm_json(system_prompt: str, human_prompt: str, response_schema) -> dict:
    """
    Calls LLM in JSON mode, validates against the response_schema.
    Performs 1 corrective retry if validation fails, falling back to gemini-2.5-flash.
    """
    model_name = settings.resolved_model
    api_key = settings.llm_api_key
    base_url = settings.resolved_base_url
    
    def run_inference(model: str, url: str, key: str, msgs) -> str:
        llm = ChatOpenAI(
            model=model,
            api_key=key,
            base_url=url,
            temperature=0.2,
            max_tokens=4096,
            timeout=120,
        ).bind(response_format={"type": "json_object"})
        resp = llm.invoke(msgs)
        return resp.content.strip()

    def clean_json_str(text: str) -> str:
        import re
        # Remove thinking blocks if present
        text = re.sub(r"<think>.*?</think>", "", text, flags=re.DOTALL | re.IGNORECASE).strip()
        # Remove code blocks if present
        if text.startswith("```"):
            lines = text.split("\n")
            text = "\n".join(lines[1:-1]) if lines[-1].strip() == "```" else "\n".join(lines[1:])
        text = text.strip()
        start = text.find("{")
        end = text.rfind("}")
        if start != -1 and end != -1 and end > start:
            return text[start:end+1]
        return text

    # Attempt 1
    messages = [
        SystemMessage(content=system_prompt),
        HumanMessage(content=human_prompt)
    ]
    
    logger.info(f"Attempting parsing with model={model_name}...")
    try:
        raw_response = run_inference(model_name, base_url, api_key, messages)
        cleaned_json = clean_json_str(raw_response)
        parsed_dict = json.loads(cleaned_json)
        validated_data = response_schema.model_validate(parsed_dict)
        return validated_data.model_dump()
    except Exception as e:
        logger.warning(f"First attempt failed: {str(e)}. Retrying with corrective prompt...")
        bad_response_str = ""
        try:
            bad_response_str = raw_response
        except NameError:
            pass

        # Try with same model but corrective re-prompt
        corrective_msgs = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=human_prompt)
        ]
        if bad_response_str:
            corrective_msgs.append(AIMessage(content=bad_response_str))
        
        validation_error_details = str(e)
        feedback_prompt = (
            f"Your previous response was invalid JSON or did not match the required schema.\n"
            f"Validation Error: {validation_error_details}\n"
            f"Please correct the output. Return ONLY the valid JSON object conforming strictly to the schema, "
            f"without any explanation, prefix, suffix, or markdown formatting."
        )
        corrective_msgs.append(HumanMessage(content=feedback_prompt))
        
        try:
            raw_response2 = run_inference(model_name, base_url, api_key, corrective_msgs)
            cleaned_json2 = clean_json_str(raw_response2)
            parsed_dict2 = json.loads(cleaned_json2)
            validated_data2 = response_schema.model_validate(parsed_dict2)
            return validated_data2.model_dump()
        except Exception as retry_err:
            logger.warning(f"Same-model corrective retry failed: {str(retry_err)}. Falling back to provider fallback model...")
            
            # Use provider-specific fallback model
            fb_model, fb_url, fb_key = _get_fallback_model_details()
            logger.info(f"Using fallback model={fb_model}...")
            
            try:
                # We can retry on the fallback model using the corrective prompt context
                raw_response3 = run_inference(fb_model, fb_url, fb_key, corrective_msgs)
                cleaned_json3 = clean_json_str(raw_response3)
                parsed_dict3 = json.loads(cleaned_json3)
                validated_data3 = response_schema.model_validate(parsed_dict3)
                return validated_data3.model_dump()
            except Exception as fb_err:
                logger.error(f"Fallback model also failed: {str(fb_err)}")
                raise HTTPException(
                    status_code=422,
                    detail=f"Failed to parse and validate JSON even after corrective retry and fallback. Details: {str(fb_err)}"
                )


@app.post("/parse-jd", response_model=JDStructuredOutput)
async def parse_jd(req: ParseJDRequest):
    """
    Parses raw JD text and returns structured data matching the JD schema.
    """
    if not req.jd_text.strip():
        raise HTTPException(status_code=400, detail="Job description text cannot be empty")
        
    jd_schema_str = json.dumps(JDStructuredOutput.model_json_schema(), indent=2)
    system_prompt = (
        "You are a precise job description analyst.\n"
        "Extract structured information from the job description.\n"
        "You MUST return ONLY valid JSON matching this schema:\n"
        f"{jd_schema_str}\n"
        "Do not return any explanations, markdown code blocks, or preamble. Return ONLY the JSON object."
    )
    human_prompt = f"Extract structured information from this job description:\n\n{req.jd_text}"
    
    parsed_data = _call_llm_json(system_prompt, human_prompt, JDStructuredOutput)
    return parsed_data


@app.post("/parse-resume", response_model=ResumeStructuredOutput)
async def parse_resume(file: UploadFile = File(...)):
    """
    Takes a resume file (PDF/DOCX/text), extracts text, and returns structured data matching the resume schema.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="Uploaded file must have a filename")
        
    file_bytes = await file.read()
    extracted_text = _extract_text_from_file(file_bytes, file.filename)
    
    if not extracted_text.strip():
        raise HTTPException(status_code=422, detail="Extracted text is empty or could not be read from file")
        
    resume_schema_str = json.dumps(ResumeStructuredOutput.model_json_schema(), indent=2)
    system_prompt = (
        "You are a precise resume parser.\n"
        "Extract structured professional information from the resume text.\n"
        "You MUST return ONLY valid JSON matching this schema:\n"
        f"{resume_schema_str}\n"
        "Do not return any explanations, markdown code blocks, or preamble. Return ONLY the JSON object."
    )
    human_prompt = f"Parse this candidate data into structured JSON:\n\n{extracted_text}"
    
    parsed_data = _call_llm_json(system_prompt, human_prompt, ResumeStructuredOutput)
    return parsed_data


# ================================================================
# ENTRYPOINT
# ================================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "api.main:app",
        host=settings.app_host,
        port=settings.app_port,
        reload=settings.debug,
    )
